


from docx import Document
import os

def add_files_to_doc(doc, folder_path):
    """
    A recursive function to go through all files and subfolders.
    """
    # List all items in the current directory
    for item_name in os.listdir(folder_path):
        # Create the full path to the item
        item_path = os.path.join(folder_path, item_name)
        
        # Check if it's a file
        if os.path.isfile(item_path):
            # Create the title in the format you want: 'folder/subfolder/file.py'
            # This gets the path relative to the original starting folder
            relative_path = os.path.relpath(item_path, start=project_path)
            
            # Add the relative path as a heading
            doc.add_heading(relative_path, level=2) # Using level 2 for slightly smaller headings
            
            # Try to read the file content as text
            try:
                with open(item_path, 'r', encoding='utf-8') as file:
                    code_content = file.read()
                # Add the code content
                doc.add_paragraph(code_content)
                
            except UnicodeDecodeError:
                # If the file can't be read as text (e.g., it's an image, .exe, etc.)
                doc.add_paragraph("(Binary file - content not displayed)")
            
            # Add a space between files
            doc.add_paragraph()
            
        # Check if it's a directory (folder)
        elif os.path.isdir(item_path):
            # Ignore the .git folder and other hidden directories
            if not item_name.startswith('.'):
                # If it's a folder, call this function again on the subfolder (RECURSION)
                add_files_to_doc(doc, item_path)

# --- MAIN SCRIPT STARTS HERE ---

# Create a new Word document
doc = Document()
# Add a main title
doc.add_heading('Project Code Documentation', level=1)

# Path to your main project folder (where you run the script from)
project_path = r'C:\Users\gopin\Documents\test_one\NeuroTrain'  


# Start the recursive process from the main project folder
add_files_to_doc(doc, project_path)

# Save the document
output_filename = 'Full_Project_Documentation.docx'
doc.save(output_filename)

print(f"Success! Document saved as '{output_filename}'")
print("It includes all files from every subfolder.")
